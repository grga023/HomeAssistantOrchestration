#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
assemble_docx.py — sklapa poglavlja (build_thesis/*.json) u profesionalni
master-rad.docx pomocu python-docx.

Kljucna razlika u odnosu na build-docx.sh: NE ugradjuje PNG slike. Sve figure iz
poglavlja se renderuju kao NATIVNI dijagrami u samom Word-u (proporcionalni
bar-dijagrami od obojenih celija tabele, box-and-arrow dijagram arhitekture i
uporedne tabele) — editabilni, konzistentnog stila, iz realnih izmerenih brojeva.

Naslovi koriste Word stilove Heading 1..5 BEZ rucne numeracije (numeracija se
dodaje naknadno). Telo: Times New Roman 12, prored 1.5, obostrano poravnanje.

Upotreba:  python assemble_docx.py [izlaz.docx]
"""
import csv
import json
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:  # noqa: BLE001
    pass

BASE = os.path.dirname(os.path.abspath(__file__))
BUILD = os.path.join(BASE, "build_thesis")

# Redosled poglavlja u finalnom dokumentu
CHAPTER_ORDER = [
    "00-sazetak",
    "01-uvod",
    "02-teorijski-okvir",
    "03-opis-sistema",
    "04-opis-komponenti",
    "05-metodologija",
    "06-rezultati",
    "07-diskusija",
    "08-provera-hipoteze",
    "09-zakljucak",
    "10-literatura",
]

# --- Paleta (uskladjena sa plot_results.py) ---------------------------------
BLUE = "2a78d6"   # plain
ORANGE = "eb6834"  # TLS
GREEN = "3fa34d"
GRAY = "9a9a94"
HDR_FILL = "1f3864"   # tamnoplava zaglavlja tabela
HDR_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
CODE_FILL = "f2f1ec"
INK = RGBColor(0x11, 0x11, 0x11)
USABLE_IN = 6.3   # upotrebljiva sirina stranice (A4, margine 2.5cm)

fig_no = 0
tbl_no = 0


# ============================================================================
# Low-level OOXML helperi
# ============================================================================
def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexc)
    tcPr.append(shd)


def _cell_valign(cell, val="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcPr.append(v)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        borders.append(e)
    tblPr.append(borders)


def _fixed_width(table, w_in):
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(w_in * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _row_height(row, twips, rule="exact"):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(twips))
    h.set(qn("w:hRule"), rule)
    trPr.append(h)


def _keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    k = OxmlElement("w:keepNext")
    pPr.append(k)


# ============================================================================
# Inline markup: **bold**, *italic*, `mono`
# ============================================================================
_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")


def add_runs(paragraph, text):
    for piece in _INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = paragraph.add_run(piece[2:-2])
            r.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = paragraph.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10.5)
        elif piece.startswith("*") and piece.endswith("*"):
            r = paragraph.add_run(piece[1:-1])
            r.italic = True
        else:
            paragraph.add_run(piece)


# ============================================================================
# Captions
# ============================================================================
def add_caption(doc, kind, text):
    global fig_no, tbl_no
    if kind == "fig":
        fig_no += 1
        label = f"Slika {fig_no}. "
    else:
        tbl_no += 1
        label = f"Tabela {tbl_no}. "
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind == "fig" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _keep_with_next(p)
    r = p.add_run(label)
    r.bold = True
    r.italic = True
    r.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(9.5)


# ============================================================================
# Native dijagrami
# ============================================================================
def bar_diagram(doc, rows, note=None):
    """rows: [(label, value, unit, hexcolor)] -> horizontalni proporcionalni barovi."""
    maxv = max((v for _, v, _, _ in rows), default=1) or 1
    for label, value, unit, color in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        rr = p.add_run(f"{label}:  ")
        rr.font.size = Pt(9.5)
        rv = p.add_run(f"{value:g} {unit}")
        rv.bold = True
        rv.font.size = Pt(9.5)
        rv.font.color.rgb = RGBColor.from_string(color)
        w = max(0.06, USABLE_IN * value / maxv)
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        _fixed_width(t, w)
        c = t.cell(0, 0)
        c.width = Inches(w)
        _shade(c, color)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        c.paragraphs[0].paragraph_format.space_before = Pt(0)
        _row_height(t.rows[0], 150, "exact")
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        r = p.add_run(note)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x52, 0x51, 0x4e)


def styled_table(doc, header, rows, first_col_bold=True, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        _shade(hdr[i], HDR_FILL)
        _cell_valign(hdr[i], "center")
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(2)
        run = para.add_run(str(h))
        run.bold = True
        run.font.color.rgb = HDR_TEXT
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.space_before = Pt(1)
            _cell_valign(cells[i], "center")
            add_runs(para, str(val))
            for rn in para.runs:
                rn.font.size = Pt(9.5)
                if first_col_bold and i == 0:
                    rn.bold = True
    return t


def value_table_from_csv(doc, path):
    if not os.path.exists(path):
        return
    with open(path, newline="", encoding="utf-8") as f:
        data = list(csv.DictReader(f))
    if not data:
        return
    header = ["Vreme [min]", "S1 [°C]", "S2 [°C]", "S3 [°C]", "S4 [°C]"]
    rows = []
    for r in data:
        mins = float(r["elapsed_s"]) / 60.0
        rows.append([f"{mins:.1f}"] + [f"{float(r[f't{i}']):.1f}" for i in range(1, 5)])
    styled_table(doc, header, rows, first_col_bold=True)


def box(cell, text, fill, textcolor=None, bold=True):
    _shade(cell, fill)
    _cell_valign(cell, "center")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(2)
    for line in text.split("\n"):
        run = para.add_run(line)
        run.bold = bold
        run.font.size = Pt(9.5)
        run.font.color.rgb = textcolor if textcolor else RGBColor(0xFF, 0xFF, 0xFF)
        run.add_break()


def arch_flow(doc):
    """Box-and-arrow dijagram arhitekture sistema (native)."""
    t = doc.add_table(rows=1, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders(t)
    cells = t.rows[0].cells
    box(cells[0], "esp1_lights\n(plain, 1883)\nESP32 WROOM", BLUE)
    box(cells[2], "Mosquitto broker\n(1883 / 8883)\nRaspberry Pi 4", "1f3864")
    box(cells[4], "Home Assistant\n(MQTT Discovery)\nHAOS", GREEN)
    for i in (1, 3):
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _cell_valign(cells[i], "center")
        r = p.add_run("→")
        r.bold = True
        r.font.size = Pt(16)
    # drugi red: esp2 ispod esp1 (MQTTS grana)
    t2 = doc.add_table(rows=1, cols=5)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders(t2)
    c2 = t2.rows[0].cells
    box(c2[0], "esp2_temperature\n(MQTTS, 8883)\nESP32 WROVER + LCD", ORANGE)
    p = c2[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("↗")
    r.bold = True
    r.font.size = Pt(16)


def sequence_flow(doc, steps):
    """Vertikalni flowchart (box + strelica nadole) — native."""
    for i, step in enumerate(steps):
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fixed_width(t, 4.6)
        c = t.cell(0, 0)
        _shade(c, "e7edf7")
        _cell_valign(c, "center")
        para = c.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(2)
        run = para.add_run(step)
        run.font.size = Pt(9.5)
        run.font.color.rgb = INK
        # border box
        tblPr = t._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement("w:" + edge)
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "6")
            e.set(qn("w:color"), "1f3864")
            borders.append(e)
        tblPr.append(borders)
        if i < len(steps) - 1:
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.paragraph_format.space_after = Pt(0)
            ap.paragraph_format.space_before = Pt(0)
            ar = ap.add_run("↓")
            ar.bold = True
            ar.font.size = Pt(13)


def code_block(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fixed_width(t, USABLE_IN)
    c = t.cell(0, 0)
    _shade(c, CODE_FILL)
    para = c.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(2)
    lines = text.split("\n")
    for j, line in enumerate(lines):
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if j < len(lines) - 1:
            run.add_break()


# ============================================================================
# Mapiranje figura -> native dijagram
# ============================================================================
def render_figure(doc, block):
    fname = os.path.basename(block.get("file", ""))
    cap = block.get("caption", "")

    if fname == "fig-heap.png":
        add_caption(doc, "fig", cap or "Zadržani heap posle povezivanja: plain vs MQTTS.")
        bar_diagram(doc, [
            ("Plain (esp1, 1883)", 7.35, "KB", BLUE),
            ("MQTTS (esp2, 8883)", 40.65, "KB", ORANGE),
        ], note="TLS trošak: +33,3 KB (≈ 15 % dostupnog heap-a posle Wi-Fi asocijacije).")
    elif fname == "fig-handshake.png":
        add_caption(doc, "fig", cap or "Vreme povezivanja / TLS handshake (median): plain vs MQTTS.")
        bar_diagram(doc, [
            ("Plain — median", 56, "ms", BLUE),
            ("MQTTS — median", 584, "ms", ORANGE),
        ], note="Jednokratni trošak po povezivanju: +528 ms (median). p95: 470 ms plain, 1048 ms MQTTS.")
    elif fname == "fig-packet.png":
        add_caption(doc, "fig", cap or "Veličina PUBLISH poruke na žici: plain vs MQTTS.")
        bar_diagram(doc, [
            ("Plain (1883)", 45, "B", BLUE),
            ("MQTTS (8883)", 74, "B", ORANGE),
        ], note="+29 B po PUBLISH-u = 5 B TLS record header + 8 B nonce + 16 B AES-GCM tag.")
    elif fname == "fig-rtt.png":
        add_caption(doc, "fig", cap or "End-to-end command RTT (p95) sa granicom hipoteze od 200 ms.")
        bar_diagram(doc, [
            ("200 ms — granica hipoteze", 200, "ms", GRAY),
            ("Plain — p95", 182.11, "ms", BLUE),
            ("TLS — p95", 189.60, "ms", ORANGE),
        ], note="Oba kanala su ispod 200 ms; mean-delta je samo +1,2 ms (TLS je praktično besplatan po komandi).")
    elif fname == "fig-cpu.png":
        add_caption(doc, "fig", cap or "Iskorišćenost procesora po fazama (avg dva jezgra).")
        bar_diagram(doc, [
            ("esp1 — povezivanje", 28, "%", BLUE),
            ("esp2 — povezivanje (TLS)", 34, "%", ORANGE),
            ("oba — ustaljeno stanje", 1, "%", GRAY),
            ("esp1 — load test (~29 kom/s)", 14, "%", GREEN),
        ], note="TLS CPU trošak je tranzijentan (+6 % u prozoru povezivanja); ustaljeno stanje ~1 %.")
    elif fname == "fig-timeseries.png":
        add_caption(doc, "tbl", cap or "esp2: četiri simulirana temperaturna senzora kroz vreme (uživo, v3.0.0).")
        value_table_from_csv(doc, os.path.join(BASE, "data", "temp_timeseries.csv"))
    elif fname == "ws-serverhello.png":
        add_caption(doc, "tbl", cap or "TLS parametri dogovoreni u handshake-u (Wireshark ServerHello + broker log).")
        styled_table(doc, ["Polje", "Vrednost"], [
            ["TLS verzija", "TLSv1.2"],
            ["Cipher suite", "`ECDHE-RSA-AES256-GCM-SHA384`"],
            ["Razmena ključeva", "ECDHE (forward secrecy)"],
            ["Autentifikacija servera", "RSA-2048"],
            ["Simetrična enkripcija", "AES-256-GCM (AEAD)"],
            ["Verifikacija lanca", "Verify return code: 0"],
        ])
    elif fname == "ws-stream-1883.png":
        add_caption(doc, "tbl", cap or "Dokaz bezbednosti: čitljivost saobraćaja plain (1883) vs MQTTS (8883).")
        styled_table(doc, ["Aspekt", "Plain (1883)", "MQTTS (8883)"], [
            ["Kredencijali (`homeuser` / `N7kQ2pX9`)", "čitljivi u čistom tekstu", "Neprozirni (enkriptovani)"],
            ["Topici i payload-i", "čitljivi", "Neprozirni"],
            ["Format na žici", "MQTT (cleartext)", "TLS Application Data"],
            ["Pasivno prisluškivanje", "Visok rizik — preuzimanje kontrole", "Eliminisan"],
        ])
    elif fname in ("ws-stream-8883.png", "ws-hierarchy.png", "ws-packetlist.png"):
        pass  # informacija je pokrivena uporednom tabelom / kodom
    else:
        # nepoznata figura: renderuj bar iz results.csv ako moze, inace preskoci
        pass


# ============================================================================
# Renderovanje blokova
# ============================================================================
def render_block(doc, block, chapter_id):
    typ = block.get("type")
    if typ == "h":
        level = max(1, min(5, int(block.get("level", 1))))
        if level == 1:
            doc.add_page_break()
        h = doc.add_heading(level=level)
        add_runs(h, block.get("text", ""))
        _keep_with_next(h)
        # posle H1 u poglavlju "opis sistema" ubaci dijagram arhitekture
        if level == 1 and chapter_id == "03-opis-sistema":
            p = doc.add_paragraph()
            r = p.add_run("Arhitektura sistema na najvišem nivou apstrakcije:")
            r.italic = True
            r.font.size = Pt(9.5)
            arch_flow(doc)
            add_caption(doc, "fig", "Topologija sistema: dva ESP32 čvora, MQTT broker i Home Assistant.")
    elif typ == "p":
        p = doc.add_paragraph()
        add_runs(p, block.get("text", ""))
    elif typ == "list":
        style = "List Number" if block.get("ordered") else "List Bullet"
        for item in block.get("items", []):
            p = doc.add_paragraph(style=style)
            add_runs(p, str(item))
    elif typ == "table":
        cap = block.get("caption")
        if cap:
            add_caption(doc, "tbl", cap)
        styled_table(doc, block.get("header", []), block.get("rows", []))
    elif typ == "figure":
        render_figure(doc, block)
    elif typ == "code":
        code_block(doc, block.get("text", ""))
    elif typ == "quote":
        p = doc.add_paragraph(style="Intense Quote")
        add_runs(p, block.get("text", ""))


# ============================================================================
# Stilizacija dokumenta
# ============================================================================
def _set_lang(style, lang="sr-Latn-RS"):
    rpr = style.element.get_or_add_rPr()
    langel = rpr.find(qn("w:lang"))
    if langel is None:
        langel = OxmlElement("w:lang")
        rpr.append(langel)
    langel.set(qn("w:val"), lang)


def style_document(doc):
    # A4 + margine
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    _set_lang(normal)

    sizes = {1: 16, 2: 14, 3: 12.5, 4: 12, 5: 11}
    for lvl, sz in sizes.items():
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Times New Roman"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1f, 0x38, 0x64)
        st.paragraph_format.space_before = Pt(12 if lvl <= 2 else 8)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_lang(st)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); run._r.append(fc)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = r'TOC \o "1-3" \h \z \u'; run._r.append(it)
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "separate"); run._r.append(fc)
    t = OxmlElement("w:t")
    t.text = "Sadržaj — desni klik → Update Field (F9) za ažuriranje."
    run._r.append(t)
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "end"); run._r.append(fc)


def title_page(doc):
    def center(text, size, bold=False, before=0, after=6, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        return p

    center("UNIVERZITET — [naziv univerziteta]", 12, bold=True, before=6)
    center("[naziv fakulteta]", 12)
    center("[katedra / smer]", 11, after=48)
    center("MASTER RAD", 13, bold=True, after=24)
    center("Analiza implementacije enkriptovanog MQTT protokola (MQTTS)",
           18, bold=True, before=12, after=2)
    center("u Smart Home sistemu", 18, bold=True, after=40)
    center("Eksperimentalno merenje troška TLS sloja na ESP32 čvorovima "
           "— memorija, vreme handshake-a, mrežni overhead, odziv i CPU",
           11, italic=True, after=48)

    # kandidat / mentor
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.cell(0, 0).text = "Kandidat:"
    t.cell(0, 1).text = "Ognjen Grgur"
    t.cell(1, 0).text = "Mentor:"
    t.cell(1, 1).text = "[ime i prezime mentora]"
    for row in t.rows:
        for i, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    r.font.size = Pt(11)
                    r.bold = (i == 0)
    center("", 11, after=36)
    center("[grad], 2026.", 11)


# ============================================================================
# Glavni tok
# ============================================================================
def load_chapter(cid):
    path = os.path.join(BUILD, f"{cid}.json")
    if not os.path.exists(path):
        return None
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:  # noqa: BLE001
        print(f"[!] {cid}: nevalidan JSON ({e})", file=sys.stderr)
        return None


def main():
    out = sys.argv[1] if len(sys.argv) > 1 else os.path.join(BASE, "master-rad.docx")
    doc = Document()
    style_document(doc)

    title_page(doc)
    doc.add_page_break()
    h = doc.add_heading("Sadržaj", level=1)
    _keep_with_next(h)
    add_toc(doc)

    present, missing = [], []
    for cid in CHAPTER_ORDER:
        ch = load_chapter(cid)
        if not ch:
            missing.append(cid)
            continue
        present.append(cid)
        for block in ch.get("blocks", []):
            render_block(doc, block, cid)

    doc.save(out)

    # --- Verifikacija ---
    d2 = Document(out)
    heads = [(int(p.style.name.split()[-1]), p.text)
             for p in d2.paragraphs
             if p.style.name.startswith("Heading") and p.style.name.split()[-1].isdigit()]
    print(f"OK: {out} ({os.path.getsize(out)//1024} KB)")
    print(f"Poglavlja: {len(present)}/{len(CHAPTER_ORDER)} prisutno; nedostaje: {missing or '-'}")
    print(f"Tabela: {len(d2.tables)} | H1: {sum(1 for l,_ in heads if l==1)} | ukupno naslova: {len(heads)}")
    print(f"Slika: {fig_no} | numerisanih tabela (caption): {tbl_no}")
    print("--- STRUKTURA (H1/H2) ---")
    for lvl, txt in heads:
        if lvl <= 2 and txt.strip():
            print(("  " * (lvl - 1)) + f"{'#'*lvl} {txt}")


if __name__ == "__main__":
    main()
